# --- IMPORTS ---
import streamlit as st
import requests
import openai
from bs4 import BeautifulSoup
import datetime
from docx import Document
from io import BytesIO

# --- SET PAGE CONFIG (FIRST Streamlit Command) ---
st.set_page_config(page_title="Pehli BackChodi - LinkedIn Post Generator - ET and TOI", layout="wide")

# --- PASSWORD PROTECTION ---
def check_password():
    st.markdown("<h3 style='text-align: center; color: #0e1117;'>🔒 Pehli BackChodi - LinkedIn Post Generator - ET and TOI</h3>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #555;'>Please enter your password to continue</p>", unsafe_allow_html=True)

    def password_entered():
        if st.session_state["password"] == "Qwerty@123!!!":
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # don't store password
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input("🔑 Password", type="password", key="password")
        if st.button("✅ Submit Password"):
            password_entered()
        return False
    elif not st.session_state["password_correct"]:
        st.text_input("🔑 Password", type="password", key="password")
        if st.button("✅ Submit Password"):
            password_entered()
        st.error("❌ Incorrect Password")
        return False
    else:
        return True

if not check_password():
    st.stop()

# --- OPENAI API KEY ---
openai.api_key = st.secrets["openai"]["api_key"]

# --- STREAMLIT PAGE SETUP ---
st.markdown(
    """
    <style>
    body {
        background-color: #f0f4f8;
        font-family: 'Segoe UI', sans-serif;
    }
    .headline {
        font-weight: 600;
        font-size: 16px;
        color: #222;
    }
    .summary {
        font-size: 13px;
        color: #555;
    }
    .source-label {
        font-size: 12px;
        color: #888;
    }
    </style>
    """, unsafe_allow_html=True
)

st.markdown("<h2 style='text-align: center;'>📰 Curated Financial & Policy News | By Gaurav Sharma</h2>", unsafe_allow_html=True)
st.caption(f"Updated on: {datetime.datetime.now().strftime('%A, %d %B %Y')}")

# --- FILTER KEYWORDS ---
FILTER_KEYWORDS = ["tax", "policy", "finance", "budget", "economy", "geopolitics", "politics"]

# --- FETCH ECONOMIC TIMES NEWS ---
def fetch_economic_times_news():
    url = "https://economictimes.indiatimes.com/news/economy"
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.content, "html.parser")

    headlines = []
    links = soup.find_all('a', href=True)

    for link in links:
        text = link.get_text(strip=True)
        href = link['href']
        
        # Focus only on economy news
        if '/news/economy/' in href and any(keyword.lower() in text.lower() for keyword in ["tax", "policy", "finance", "budget", "economy", "geopolitics", "infrastructure", "inflation"]):
            if len(text) > 10:  # basic text length check
                full_link = "https://economictimes.indiatimes.com" + href if not href.startswith('http') else href
                headlines.append(("ET", text, full_link))

    return headlines

def fetch_times_of_india_news():
    url = "https://timesofindia.indiatimes.com/india"
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.content, "html.parser")

    headlines = []
    links = soup.find_all('a', href=True)

    for link in links:
        text = link.get_text(strip=True)
        href = link['href']
        
        if '/india/' in href and any(keyword.lower() in text.lower() for keyword in ["tax", "policy", "finance", "budget", "economy", "geopolitics", "infrastructure", "inflation"]):
            if len(text) > 10:
                full_link = "https://timesofindia.indiatimes.com" + href if not href.startswith('http') else href
                headlines.append(("TOI", text, full_link))

    return headlines

# --- COMBINE NEWS ---
def get_filtered_news():
    news = fetch_economic_times_news() + fetch_times_of_india_news()
    return news[:10]

# --- SMART 3-POST LINKEDIN POST GENERATOR ---

def generate_3_linkedin_posts(news_title):
    styles = {
        "Analytical Professional Tone": """
You are a senior finance editor and economist. Write a premium LinkedIn post:

1. Start with a catchy professional headline (max 10 words).
2. Write 1-line simple summary of what happened.
3. Add a short paragraph explaining why this matters.
4. Create 3–5 bullet points analyzing:
    - Impact/Opportunities
    - Risks/Challenges
    - Global Comparisons
    - Future Outlook
5. Insert a personal reflection.
6. Finish with an intelligent, thought-provoking question.
7. Add 4-5 relevant finance/policy hashtags.

Frame numbers carefully like "as per available data" or "historical trends suggest".
""",
        
        "Bold Straight-Talk Tone": """
You are a bold and direct commentator. Write a LinkedIn post:

1. Start with a strong emotional headline (max 10 words).
2. Write a 1-line simple "What Happened".
3. Add 1 short para explaining risk or opportunity.
4. Bullet points covering:
    - What’s good/bad
    - Immediate risks
    - Short-term impact
5. Insert strong personal opinion.
6. End with a bold question.
7. Add 4-5 mass-appeal hashtags.

Frame numbers responsibly ("estimated", "historical averages").
""",
        
        "Visionary Long-Term Tone": """
You are a macro-strategist. Write a visionary LinkedIn post:

1. Catchy visionary headline (max 10 words).
2. 1-line "What Happened" summary.
3. Small para on long-term significance.
4. Bullets covering:
    - Future potential
    - Strategic risks
    - Global trends
    - Opportunities
5. Inspirational personal reflection.
6. End with a futuristic, visionary question.
7. Add 4-5 forward-looking hashtags.

Frame any numbers carefully as "historical patterns" or "trends indicate".
"""
    }

    responses = {}

    for style, prompt in styles.items():
        try:
            chat_completion = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert finance, policy and macroeconomic content strategist."},
                    {"role": "user", "content": f"News Headline: {news_title}\n\n{prompt}"}
                ],
                temperature=0.6,
                top_p=0.9
            )
            responses[style] = chat_completion.choices[0].message.content.strip()
        except Exception as e:
            responses[style] = f"(Error generating post: {e})"
    
    return responses


# --- MAIN APP LAYOUT ---

st.markdown("### 🔍 Today's Curated Financial & Policy Headlines")

news_items = get_filtered_news()

for index, (source, title, link) in enumerate(news_items):
    with st.container():
        st.markdown(f"<div class='headline'>{index+1}. {title}</div>", unsafe_allow_html=True)
        st.markdown(f"<span class='source-label'>Source: [{source}]  |  [🔗 Read more]({link})</span>", unsafe_allow_html=True)

        if st.button(f"✍️ Generate 3 Premium LinkedIn Posts", key=f"gen_{index}"):
            posts = generate_3_linkedin_posts(title)

            for style, post_text in posts.items():
                with st.expander(f"💬 {style}"):
                    st.text_area(f"Generated Post ({style})", value=post_text, height=400, key=f"ta_{index}_{style}")

                    # Choose and download
                    if st.button(f"📥 Use This {style} Post", key=f"use_{index}_{style}"):
                        buffer = BytesIO()
                        doc = Document()
                        doc.add_heading(f"LinkedIn Post - {style}", level=1)
                        doc.add_paragraph(post_text)
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            f"📄 Download {style} Post", 
                            data=buffer, 
                            file_name=f"LinkedIn_{style.replace(' ', '_')}.docx", 
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                            key=f"dl_{index}_{style}"
                        )
                        st.code(post_text, language="markdown")

# --- FOOTER ---
st.markdown("---")
st.caption("🛠 Crafted with ❤️ by Gaurav Sharma | Powered by GPT-4")

